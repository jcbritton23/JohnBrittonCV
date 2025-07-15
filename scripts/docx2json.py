import sys
import json
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_formatted_text_from_paragraph(paragraph):
    formatted_parts = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            formatted_parts.append(f"<strong><em>{text}</em></strong>")
        elif run.bold:
            formatted_parts.append(f"<strong>{text}</strong>")
        elif run.italic:
            formatted_parts.append(f"<em>{text}</em>")
        else:
            formatted_parts.append(text)
    return "".join(formatted_parts)

def extract_table_content(table):
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = ""
            for paragraph in cell.paragraphs:
                cell_text += extract_formatted_text_from_paragraph(paragraph)
            row_data.append(cell_text.strip())
        table_data.append(row_data)
    return table_data

def parse_docx_to_json(docx_path):
    doc = Document(docx_path)
    all_blocks = []
    para_idx = 0
    table_idx = 0
    # Walk through the document in order (paragraphs and tables)
    for element in doc.element.body:
        if element.tag.endswith('p'):
            paragraph = doc.paragraphs[para_idx]
            formatted_text = extract_formatted_text_from_paragraph(paragraph)
            all_blocks.append({
                "type": "paragraph",
                "text": formatted_text,
                "raw_text": paragraph.text,
                "alignment": str(paragraph.alignment)
            })
            para_idx += 1
        elif element.tag.endswith('tbl'):
            table = doc.tables[table_idx]
            table_data = extract_table_content(table)
            all_blocks.append({
                "type": "table",
                "data": table_data
            })
            table_idx += 1
    data = {
        "content_blocks": all_blocks,
        "metadata": {
            "total_blocks": len(all_blocks),
            "source_file": docx_path
        }
    }
    return data

def main():
    if len(sys.argv) != 3:
        print("Usage: python3 docx2json.py <input.docx> <output.json>")
        sys.exit(1)
    docx_path = sys.argv[1]
    json_path = sys.argv[2]
    try:
        print(f"Processing DOCX file: {docx_path}")
        data = parse_docx_to_json(docx_path)
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"Successfully converted DOCX to JSON: {json_path}")
        print(f"Extracted {len(data.get('content_blocks', []))} content blocks")
    except Exception as e:
        print(f"Error converting DOCX to JSON: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main() 